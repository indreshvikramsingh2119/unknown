
from PIL import Image, ImageTk
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib import colors
from datetime import datetime
import os
import serial.tools.list_ports
import serial
import threading
import time
from docx import Document
import pandas as pd
import customtkinter as ctk


ctk.set_appearance_mode("light")  # optional
ctk.set_default_color_theme("blue")  # optional



baud_rates = ["9600", "19200", "38400", "57600", "115200"]
def get_serial_ports():
    ports = serial.tools.list_ports.comports()
    return [port.device for port in ports]

serial_ports = get_serial_ports()

# --- Action Functions ---
def login():
    username = username_entry.get()
    password = password_entry.get()
    if username and password:
        frame.place_forget()
        canvas.pack_forget()
        show_main_page()

def show_main_page():
    root.geometry("1500x1000")
    main_frame.pack(fill="both", expand=True)

ser = None
running = False
latest_serial_values = ["", "", "", "", "", ""] 
serial_data_history = []

history_table_rows = []
history_table_frame = None



def start_serial():
    global ser, running
    selected_port = port_dropdown.get()
    selected_baud = baud_dropdown.get()

    try:
        ser = serial.Serial(selected_port, int(selected_baud), timeout=1)
        if ser.is_open:
            messagebox.showinfo("Connected", f"\ Connected to {selected_port}")
            running = True
            threading.Thread(target=read_serial_data, daemon=True).start()
    except Exception as e:
        messagebox.showerror("Error", f" Failed to connect:\n{e}")


    """
    The function `stop_serial` prints a message indicating that it has stopped reading serial data.
    """
# def stop_serial():
#     print("Stopped reading serial data...")

def stop_serial():
    global running
    running = False
    if ser.is_open:
        ser.close()
    messagebox.showinfo("Disconnected", " Serial connection closed.")





def show_history_page():
    global history_table_rows, history_table_frame
    main_frame.pack_forget()
    if print_frame:
        print_frame.pack_forget()
    history_frame.pack(fill="both", expand=True)

    # Clear old widgets
    for widget in history_frame.winfo_children():
        widget.destroy()
    history_table_rows.clear()

    # --- Canvas + Scroll for Table Only ---
    canvas_frame = tk.Frame(history_frame, bg="white")
    canvas_frame.pack(fill="both", expand=True, padx=100, pady=(20, 0))

    canvas = tk.Canvas(canvas_frame, bg="white", highlightthickness=0)
    scrollbar = tk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=scrollbar.set)

    scrollbar.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    scroll_inner_frame = tk.Frame(canvas, bg="white")
    canvas.create_window((0, 0), window=scroll_inner_frame, anchor="nw")

    def on_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))
    scroll_inner_frame.bind("<Configure>", on_configure)

    # --- Table ---
    history_table_frame = tk.Frame(scroll_inner_frame, bg="white", bd=2, relief="ridge", highlightbackground="#ccc", highlightthickness=1)
    history_table_frame.pack(fill="x")

    # headers = ["[✓]", "Date", "Time", "Initial Temp", "Max Initial Temp", "Net Rise Temp", "Benzoic Acid CV", "Water Eq.", "Calo Val"]
    headers = ["[✓]", "Date", "Time","Type", "A", "B", "C", "D", "E", "F"]

    for i, h in enumerate(headers):
        label = tk.Label(history_table_frame, text=h, font=("Segoe UI", 13, "bold"), bg="white", fg="#37474f", anchor="center")
        label.grid(row=0, column=i, padx=20, pady=12)

    for row_num, row in enumerate(serial_data_history, start=1):
        now = datetime.now()
        date = now.strftime("%Y-%m-%d")
        time_str = now.strftime("%I:%M %p")
        full_row = [date, time_str] + row

        var = tk.BooleanVar()
        chk = tk.Checkbutton(history_table_frame, variable=var, bg="skyblue")
        chk.var = var  # attach variable for later access
        chk.grid(row=row_num, column=0, padx=10)

        for col_num, value in enumerate(full_row):
            data_label = tk.Label(history_table_frame, text=value, font=("Segoe UI", 11), bg="white", anchor="center")
            data_label.grid(row=row_num, column=col_num + 1, padx=15, pady=8)

        history_table_rows.append((chk, *full_row))

    # --- Separator ---
    separator2 = tk.Frame(history_frame, bg="#90caf9", height=2)
    separator2.pack(fill="x", padx=100, pady=(10, 0))

    # --- Button Frame ---
    btn_frame = tk.Frame(history_frame, bg="#e3f2fd")
    btn_frame.pack(pady=20)

    # --- Save PDF Function ---
    def save_selected_rows_as_pdf():
        selected_data = []
        for row in history_table_rows:
            checkbox = row[0]
            if checkbox.var.get():  # if checkbox is selected
                selected_data.append(row[1:])  # skip checkbox

        if not selected_data:
            messagebox.showinfo("No Selection", "Please select at least one row to save as PDF.")
            return

        file_name = f"Selected_History_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)

        c = pdf_canvas.Canvas(file_path, pagesize=A4)
        width, height = A4
        y = height - 50

        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, "( History Data")
        y -= 30

        c.setFont("Helvetica-Bold", 12)
        # headers_to_print = ["Date", "Time", "Initial Temp", "Max Initial Temp", "Net Rise Temp", "Benzoic Acid CV", "Water Eq.", "Calo Val"]
        headers_to_print = ["Date", "Time", "A", "B", "C", "D", "E.", "F"]

        for i, h in enumerate(headers_to_print):
            c.drawString(50 + i * 65, y, h)
        y -= 20

        c.setFont("Helvetica", 11)
        for row in selected_data:
            for i, val in enumerate(row):
                c.drawString(50 + i * 65, y, str(val))
            y -= 20
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 11)

        c.save()
        messagebox.showinfo("PDF Saved", f"PDF saved to:\n{file_path}")

    # --- Save Button ---
    save_pdf_btn = tk.Button(btn_frame, text=" print Pdf", font=("Segoe UI", 12, "bold"), bg="#6a1b9a", fg="white",
                             command=save_selected_rows_as_pdf, cursor="hand2", bd=0, relief="flat")
    save_pdf_btn.pack(side="left", padx=10, ipadx=15, ipady=7)

    # --- Back to Dashboard Button ---
    back_btn = tk.Button(btn_frame, text="Back ", font=("Segoe UI", 12, "bold"), bg="#0288d1", fg="white",
                         command=back_to_main, cursor="hand2", bd=0, relief="flat")
    back_btn.pack(side="left", padx=10, ipadx=15, ipady=7)


# print_frame = None

# def show_print_page():
#     global print_frame
#     main_frame.pack_forget()
#     history_frame.pack_forget()

#     if print_frame:
#         print_frame.destroy()

#     print_frame = tk.Frame(root, bg="#f8f9fa")
#     print_frame.pack(fill="both", expand=True)
    
    
#     # now = datetime.now()
#     # date_str = now.strftime(" %d-%b-%Y")
#     # time_str = now.strftime(" %I:%M %p")

#     # datetime_frame = tk.Frame(print_frame, bg="#f8f9fa")
#     # datetime_frame.pack(anchor="ne", padx=20, pady=(10, 0))

#     # tk.Label(datetime_frame, text=date_str, font=("Segoe UI", 10, "bold"),
#     #          bg="#f8f9fa", fg="#333").pack(anchor="e")
#     # tk.Label(datetime_frame, text=time_str, font=("Segoe UI", 10),
#     #          bg="#f8f9fa", fg="#555").pack(anchor="e")
    
#     # New Top Title
#     title = tk.Label(print_frame, text=" Deckmount Report Form",
#                     font=("Segoe UI", 18, "bold"), bg="#f8f9fa", fg="#004d40")
#     title.pack(pady=(20, 5))


#     # heading = tk.Label(print_frame, text=" Fill Report Details",
#     #                    font=("Segoe UI", 16, "bold"), bg="#f8f9fa", fg="#1a1a1a")
#     # heading.pack(pady=(20, 10))
    
#     # header_frame = tk.Frame(print_frame, bg="#f8f9fa")
#     # header_frame.pack(fill="x", padx=30)

#     # heading_label = tk.Label(header_frame, text="Fill Report Details",
#     #                          font=("Segoe UI", 14, "bold"), bg="#f8f9fa", fg="#1a1a1a")
#     # heading_label.pack(side="left")

#     # now = datetime.now()
#     # date_label = tk.Label(header_frame, text=now.strftime(" %d-%b-%Y"),
#     #                       font=("Segoe UI", 10), bg="#f8f9fa", fg="#333")
#     # time_label = tk.Label(header_frame, text=now.strftime(" %I:%M %p"),
#     #                       font=("Segoe UI", 10), bg="#f8f9fa", fg="#333")
#     # date_label.pack(side="right", anchor="ne")
#     # time_label.pack(side="right", anchor="ne")

#     # -- Header Frame (Centered Heading + Date/Time in Top Right stacked) --
#     header_frame = tk.Frame(print_frame, bg="#f8f9fa")
#     header_frame.pack(fill="x", padx=30)

#     # Centered Heading
#     heading_label = tk.Label(header_frame, text="Fill Report Details",
#                             font=("Segoe UI", 14, "bold"), bg="#f8f9fa", fg="#1a1a1a")
#     heading_label.pack(side="top", pady=(10, 5))

#     # Date and Time stacked on top right
#     datetime_frame = tk.Frame(header_frame, bg="#f8f9fa")
#     datetime_frame.pack(side="right", anchor="ne")

#     now = datetime.now()

#     date_label = tk.Label(datetime_frame, text=now.strftime(" %d-%b-%Y"),
#                         font=("Segoe UI", 10), bg="#f8f9fa", fg="#333")
#     time_label = tk.Label(datetime_frame, text=now.strftime(" %I:%M %p"),
#                         font=("Segoe UI", 10), bg="#f8f9fa", fg="#333")

#     date_label.pack(anchor="e")
#     time_label.pack(anchor="e")


#     form = tk.Frame(print_frame, bg="#e8f5e9", bd=2, relief="groove", width=500)
#     form.pack(padx=30, pady=20, fill="both", expand=True)
#     # form.place(relx=0.5, rely=0.5, anchor="center")


#     fields = [
#         ("Company Name:", ""),
#         ("Company Address:", ""),
#         ("Customer ID:", ""),
#         ("Mobile No:", ""),
#         # ("Initial Temp:", ""),
#         # ("Max Initial Temp:", ""),
#         # ("Net Rise Temp:", ""),
#         # ("Benzioc Acid CV:", ""),
#         # ("Water Equivalent:", ""),
#         # ("Caloric Value:", ""),
        
#         # ("Initial Temp:", latest_serial_values[0]),
#         # ("Max Initial Temp:", latest_serial_values[1]),``
#         # ("Net Rise Temp:", latest_serial_values[2]),
#         # ("Benzioc Acid CV:", latest_serial_values[3]),
#         # ("Water Equivalent:", latest_serial_values[4]),
#         # ("Caloric Value:", latest_serial_values[5]),
#         # ]
    
        
#         ("A:", latest_serial_values[0]),
#         ("B:", latest_serial_values[1]),
#         ("C:", latest_serial_values[2]),
#         ("D:", latest_serial_values[3]),
#         ("E:", latest_serial_values[4]),
#         ("F:", latest_serial_values[5]),
#         ]
    
    
    
#     def validate_mobile(P):
#      return P == "" or (P.isdigit() and len(P) <= 10)

    
#     vcmd = (root.register(validate_mobile), '%P')
    
    
#     # unit_labels = {
#     #     "Initial Temp:": "°C",
#     #     "Max Initial Temp:": "°C",
#     #     "Net Rise Temp:": "°C",
#     #     "Benzioc Acid CV:": "cal/gram",
#     #     "Water Equivalent:": "cal/°C",
#     #     "Caloric Value:": "cal/gram"
#     # }
    
#     unit_labels = {
#         "A:": "°C",
#         "B:": "°C",
#         "C:": "°C",
#         "D:": "cal/gram",
#         "E:": "cal/°C",
#         "F:": "cal/gram"
#     }
        
#     entries = {}
#     for i, (label_text, default) in enumerate(fields):
#         tk.Label(form, text=label_text, font=("Segoe UI", 11, "bold"),
#                  bg="white").grid(row=i, column=0, padx=15, pady=8, sticky="e")

#         if label_text == "Mobile No:":
#             e = tk.Entry(form, font=("Segoe UI", 11), width=15,
#                          validate="key", validatecommand=vcmd)
#         else:
#             e = tk.Entry(form, font=("Segoe UI", 11), width=15)

#         e.insert(0, default)
#         e.grid(row=i, column=1, padx=15, pady=8, sticky="w")
#         entries[label_text] = e
        
        
#         unit = unit_labels.get(label_text)
#         if unit:
#             tk.Label(form, text=unit, font=("Segoe UI", 11,"bold"),
#                      bg="white", fg="#555").grid(row=i, column=2, padx=(5, 10), sticky="w")

    
    
#     # Button Row
#     button_row = tk.Frame(print_frame, bg="#f8f9fa")
#     button_row.pack(pady=(10, 20))

#     back_btn = tk.Button(button_row, text="⬅ Back", font=("Segoe UI", 12, "bold"),
#                          bg="#694a54", fg="white", bd=0, cursor="hand2",
#                          command=back_to_main, width=18)
#     back_btn.grid(row=0, column=0, padx=10, ipadx=5, ipady=6)   


print_frame = None

def show_print_page():
    global print_frame
    main_frame.pack_forget()
    history_frame.pack_forget()

    if print_frame:
        print_frame.destroy()

    print_frame = tk.Frame(root, bg="#f8f9fa")
    print_frame.pack(fill="both", expand=True)
    
    
    
    # New Top Title
    title = tk.Label(print_frame, text="Deckmount Report Form",
                    font=("Segoe UI",18, "bold"), bg="#f8f9fa", fg="#004d40")
    title.pack(pady=(20, 5))
    
    
    
    # -- Header Frame (Centered Heading + Date/Time in Top Right stacked) --
    header_frame = tk.Frame(print_frame, bg="#f8f9fa")
    header_frame.pack(fill="x", padx=30, pady=(10, 0))

    # Centered Heading
    heading_label = tk.Label(header_frame, text="Fill Report Details",
                            font=("Segoe UI", 14, "bold"), bg="#f8f9fa", fg="#1a1a1a")
    heading_label.pack(side="top", pady=( 5))

    # Date and Time stacked on top right
    # datetime_frame = tk.Frame(header_frame, bg="#f8f9fa")
    datetime_frame = tk.Frame(print_frame, bg="#f8f9fa")  #  use print_frame here

    # datetime_frame.pack(side="right", anchor="ne", pady=(0, 10))
    datetime_frame.place(relx=1.0, y=10, anchor="ne")  #  absolute top-right position


    now = datetime.now()

    date_label = tk.Label(datetime_frame, text=now.strftime(" %d-%b-%Y"),
                        font=("Segoe UI", 10), bg="#f8f9fa", fg="#333")
    time_label = tk.Label(datetime_frame, text=now.strftime(" %I:%M %p"),
                        font=("Segoe UI", 10), bg="#f8f9fa", fg="#333")

    date_label.pack(anchor="e")
    time_label.pack(anchor="e")

    content_frame = tk.Frame(print_frame, bg="#f8f9fa")
    content_frame.pack(padx=30, pady=20, fill="both", expand=True)
    
    total_width = 1500
    frame_width = total_width // 2
    
    
    
    
    # --- Left Frame ---
    left_frame = tk.Frame(content_frame, bg="#e0ffff", width=frame_width, bd=2, relief="groove")
    left_frame.pack(side="left", fill="both", expand=False, padx=10, pady=10)
    left_frame.pack_propagate(False)

    # --- Center Frame ---
    center_frame = tk.Frame(content_frame, bg="#f5f5dc", width=frame_width, bd=2, relief="groove")
    center_frame.pack(side="left", fill="both", expand=False, padx=10, pady=10)
    center_frame.pack_propagate(False)

    # --- Right Frame ---
    right_frame = tk.Frame(content_frame, bg="#eeeeee", width=frame_width,  )
    right_frame.pack(side="left", fill="both", expand=False, padx=10, pady=10)
    right_frame.pack_propagate(False)
    
    
    
    

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    
    
    
    fields = [
        ("Company Name:", ""),
        ("Company Address:", ""),
        ("Customer ID:", ""),
        ("Mobile No:", ""),
        # ("Initial Temp:", ""),
        # ("Max Initial Temp:", ""),
        # ("Net Rise Temp:", ""),
        # ("Benzioc Acid CV:", ""),
        # ("Water Equivalent:", ""),
        # ("Caloric Value:", ""),
        
        # ("Initial Temp:", latest_serial_values[0]),
        # ("Max Initial Temp:", latest_serial_values[1]),``
        # ("Net Rise Temp:", latest_serial_values[2]),
        # ("Benzioc Acid CV:", latest_serial_values[3]),
        # ("Water Equivalent:", latest_serial_values[4]),
        # ("Caloric Value:", latest_serial_values[5]),
        # ]
    
        
        ("A:", latest_serial_values[0]),
        ("B:", latest_serial_values[1]),
        ("C:", latest_serial_values[2]),
        ("D:", latest_serial_values[3]),
        ("E:", latest_serial_values[4]),
        ("F:", latest_serial_values[5]),
        ]
    
    
    
    def validate_mobile(P):
     return P == "" or (P.isdigit() and len(P) <= 10)

    
    vcmd = (root.register(validate_mobile), '%P')
    
    
    # unit_labels = {
    #     "Initial Temp:": "°C",
    #     "Max Initial Temp:": "°C",
    #     "Net Rise Temp:": "°C",
    #     "Benzioc Acid CV:": "cal/gram",
    #     "Water Equivalent:": "cal/°C",
    #     "Caloric Value:": "cal/gram"
    # }
    
    unit_labels = {
        "A:": "°C",
        "B:": "°C",
        "C:": "°C",
        "D:": "cal/gram",
        "E:": "cal/°C",
        "F:": "cal/gram"
    }
        
    entries = {}
    
    for i, (label_text, default) in enumerate(fields):
        # tk.Label(left_frame, text=label_text, font=("Segoe UI", 11, "bold"),
        #         bg="#f5fffa").grid(row=i, column=0, padx=15, pady=8, sticky="e")
        
        ctk.CTkLabel(left_frame, text=label_text, font=("Segoe UI", 16, "bold"),
            fg_color="#D0F0C0", corner_radius=15, width=160, height=30).grid(row=i, column=0, padx=15, pady=8, sticky="e")

        
        
        if label_text == "Mobile No:":
            e = tk.Entry(left_frame, font=("Segoe UI", 11), width=15,
                        validate="key",validatecommand=vcmd)
        else:
            e = tk.Entry(left_frame, font=("Segoe UI", 11), width=15)

        e.insert(0, default)
        e.grid(row=i, column=1, padx=15, pady=8, sticky="w")
        entries[label_text] = e
        
        
        unit = unit_labels.get(label_text)
        if unit:
            tk.Label(left_frame, text=unit, font=("Segoe UI", 11,"bold"),
                    bg="white", fg="#555").grid(row=i, column=2, padx=(5, 10), sticky="w")
            
            
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    
    
    

# --- BOX 1: 5-digit Section in White Frame ---
    


    box1 = tk.Frame(right_frame, bg="#B0E0E6", bd=2, relief="groove", padx=30, pady=40)
    box1.pack(fill="x", pady=(0, 5))
    
    
    
    # heading2 = tk.Label(box1, text="Deckmount udyog vihar", font=("Segoe UI", 22, "bold"),
    #                     bg="#B0E0E6", fg="#1a1a1a")
    # heading2.pack(pady=(0, 5))
    
    rounded_heading = ctk.CTkLabel(
        master=box1,
        text="Deckmount udyog vihar",
        font=("Segoe UI", 22, "bold"),
        fg_color="#FFCCBC",       
        text_color="#1a1a1a",     # Text color
        corner_radius=20,         # Rounded
        width=300,
        height=50
    )
    rounded_heading.pack(pady=(10, 10))
     
    # Name Section inside box1
    name_section = tk.Frame(box1, bg="white")
    name_section.pack(expand=True,pady=(25, 0))

    # Heading
    # name_label = tk.Label(name_section, text="Deckmount udyog vihar", font=("Segoe UI", 12, "bold"),
    #                     bg="#f5fffa", fg="#1a1a1a")
    # name_label.pack(pady=(0, 5))
    # heading2 = tk.Label(box1, text="Manual Fill Section", font=("Segoe UI", 12, "bold"),
    #                     bg="#f5fffa", fg="#1a1a1a")
    # heading2.pack(pady=(10, 5))

    # Validate only 5-digit input
    def validate_five_digits(P):
        return P == "" or (P.isdigit() and len(P) <= 5)

    vcmd = root.register(validate_five_digits)

    # Entry Box
    name_entry = tk.Entry(name_section, font=("Segoe UI", 12), width=10, justify="center",
                        validate="key", validatecommand=(vcmd, "%P"))
    name_entry.pack()



    box2 = tk.Frame(right_frame, bg="#B0E0E6", bd=2, relief="groove",pady= 40)
    box2.pack(side="bottom", fill="x", padx=0, pady=(50, 0), anchor="s")
    
    
    # Heading
    # heading2 = tk.Label(box2, text="Manual Fill Section", font=("Segoe UI", 22, "bold"),
    #                     bg="#B0E0E6", fg="#1a1a1a")
    # heading2.pack(pady=(0, 5))
    rounded_manual_label = ctk.CTkLabel(
        master=box2,
        text="Manual Fill Section",
        font=("Segoe UI", 22, "bold"),
        fg_color="#E1BEE7",      
        text_color="#1a1a1a",    # Text color
        corner_radius=20,        # Rounded corner
        width=300,
        height=50
    )
    rounded_manual_label.pack(pady=(10, 10))

    # Internal entries
    tk.Entry(box2, font=("Segoe UI", 12), width=25).pack(pady=10,)

    
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    
    
    auto_fields = [
        ("CV of thread:",latest_serial_values [0]),
        ("CV of unbrunt wire:",latest_serial_values[1] ),
        ("NA CV:",latest_serial_values [2]),
        ("SA CV:",latest_serial_values [3]),
        ("Other connection:",latest_serial_values [4]),
        ("Total CV:",latest_serial_values [5])
    ]
    unit_labels = {
        "NA CV:": "cal",
        "SA CV:": "cal",
        "Other connection:": "cal",
        "Total CV:": "cal"
    }



    for i, (label_text, default) in enumerate(auto_fields, start=2):
        # tk.Label(center_frame, text=label_text, font=("Segoe UI", 11, "bold"),
        #         bg="#f5fffa").grid(row=i, column=0, padx=15, pady=8, sticky="e")
        ctk.CTkLabel(center_frame, text=label_text, font=("Segoe UI", 16, "bold"),
             fg_color="#D0F0C0", corner_radius=15, width=160, height=30).grid(row=i, column=0, padx=15, pady=8, sticky="e")

        e = tk.Entry(center_frame, font=("Segoe UI", 11), width=15)
        e.insert(0, default)
        e.grid(row=i, column=1, padx=10, pady=8, sticky="w")
        entries[label_text] = e

        unit = unit_labels.get(label_text)
        if unit:
            tk.Label(center_frame, text=unit, font=("Segoe UI", 11, "bold"),
                    bg="white", fg="#555").grid(row=i, column=2, padx=(5, 10), sticky="w")

    
    
    # Button Row
    button_row = tk.Frame(print_frame, bg="#f8f9fa")
    button_row.pack(pady=(0, 10))

    back_btn = tk.Button(button_row, text="⬅ Back", font=("Segoe UI", 12, "bold"),
                         bg="#694a54", fg="white", bd=0, cursor="hand2",
                         command=back_to_main, width=18)
    back_btn.grid(row=0, column=0, padx=10, ipadx=5, ipady=6)   

   

#     # -------- PDF REPORT GENERATION -------- #
#     def generate_report():
#         file_name = f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
#         file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)

#         c = pdf_canvas.Canvas(file_path, pagesize=A4)
#         width, height = A4
#         y = height - 50

#         # === Title Bar ===
#         # c.setFillColorRGB(224/255, 247/255, 250/255)  # Light cyan blue
#         # c.rect(0, 0, width, height, fill=1)
#         c.setFillColorRGB(0.95, 0.95, 0.95)  
#         c.rect(0, height -70,width ,30,stroke=0,fill=1)

#         # Title Text
#         c.setFont("Helvetica-Bold", 16)
#         c.setFillColor(colors.HexColor("#1a237e"))
#         c.drawCentredString(width / 2, y, " PATIENT REPORT")
#         y -= 35

#         # Line
#         c.setStrokeColor(colors.lightgrey)
#         c.line(40, y, width - 40, y)
#         y -= 20

#         # Helper Function
#         def draw_section(title, lines):
#             nonlocal y
#             c.setFont("Helvetica-Bold", 12)
#             c.setFillColor(colors.HexColor("#3f51b5"))
#             c.drawString(50, y, title)
#             y -= 20

#             c.setFont("Courier", 11)
#             c.setFillColor(colors.black)
#             for line in lines:
#                 c.drawString(60, y, line)
#                 y -= 16
#                 if y < 50:
#                     c.showPage()
#                     y = height - 50

#         # Content Sections
#         draw_section(" Company Details", [
#             f"Company Name     : {entries['Company Name:'].get()}",
#             f"Company Address  : {entries['Company Address:'].get()}",
#             f"Customer ID      : {entries['Customer ID:'].get()}",
#             f"Mobile No        : {entries['Mobile No:'].get()}",
#         ])

#         # draw_section(" Temperature Analysis", [
#         #     f"Initial Temp     : {entries['Initial Temp:'].get()} °C",
#         #     f"Max Initial Temp : {entries['Max Initial Temp:'].get()} °C",
#         #     f"Net Rise Temp    : {entries['Net Rise Temp:'].get()} °C",
#         # ])

#         # draw_section(" Code Readings", [
#         #     f"Benzoic Acid CV                : {entries['Benzoic Acid CV:'].get()}",
#         #     f"Water Equivalent               : {entries['Water Equivalent:'].get()}",
#         #     f"Caloric Value                : {entries['Caloric Value:'].get()}",
#         # ])

#         draw_section(" Temperature Analysis", [
#             f"A     : {entries['A:'].get()} °C",
#             f"B : {entries['B:'].get()} °C",
#             f"C    : {entries['C:'].get()} °C",
#         ])

#         draw_section(" Code Readings", [
#             f"D                : {entries['D:'].get()}",
#             f"E               : {entries['D:'].get()}",
#             f"F                : {entries['F:'].get()}",
#         ])
#         # Footer
#         y -= 10
#         c.setFillColor(colors.gray)
#         c.drawString(50, y, "-" * 70)
#         y -= 16
#         c.setFont("Helvetica-Oblique", 10)
#         c.drawString(50, y, "-- End of Report --")

#         c.save()
#         messagebox.showinfo("PDF Saved", f"Report saved to:\n{file_path}")
        
        
#     def generate_report_excel():
#         data = {
#             "Company Name": [entries['Company Name:'].get()],
#             "Company Address": [entries['Company Address:'].get()],
#             "Customer ID": [entries['Customer ID:'].get()],
#             "Mobile No": [entries['Mobile No:'].get()],
#             "A": [entries['A:'].get()],
#             "B": [entries['B:'].get()],
#             "C": [entries['C:'].get()],
#             "D": [entries['D:'].get()],
#             "E": [entries['E:'].get()],
#             "F": [entries['F:'].get()],
#         }

#         df = pd.DataFrame(data)
#         file_name = f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
#         file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)
#         df.to_excel(file_path, index=False)
#         messagebox.showinfo("Excel Saved", f"Report saved to:\n{file_path}")


# def generate_report_word():
#     doc = Document()
#     doc.add_heading('Patient Report', 0)

#     doc.add_heading('Company Details', level=1)
#     doc.add_paragraph(f"Company Name: {entries['Company Name:'].get()}")
#     doc.add_paragraph(f"Company Address: {entries['Company Address:'].get()}")
#     doc.add_paragraph(f"Customer ID: {entries['Customer ID:'].get()}")
#     doc.add_paragraph(f"Mobile No: {entries['Mobile No:'].get()}")

#     doc.add_heading('Temperature Analysis', level=1)
#     doc.add_paragraph(f"A: {entries['A:'].get()} °C")
#     doc.add_paragraph(f"B: {entries['B:'].get()} °C")
#     doc.add_paragraph(f"C: {entries['C:'].get()} °C")

#     doc.add_heading('Code Readings', level=1)
#     doc.add_paragraph(f"D: {entries['D:'].get()}")
#     doc.add_paragraph(f"E: {entries['E:'].get()}")
#     doc.add_paragraph(f"F: {entries['F:'].get()}")

#     file_name = f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
#     file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)
#     doc.save(file_path)
#     messagebox.showinfo("Word Saved", f"Report saved to:\n{file_path}")


#     # generate_btn = tk.Button(button_row, text=" Generate Report", font=("Segoe UI", 12, "bold"),
#     #                           bg="#1a759f", fg="white", bd=0, cursor="hand2",
#     #                           command=generate_report, width=18)
#     # generate_btn.grid(row=0, column=1, padx=10, ipadx=5, ipady=6)
    
#     def save_report_format(choice):
#         if choice == "Save as PDF":
#             generate_report_pdf()
#         elif choice == "Save as Excel":
#             generate_report_excel()
#         elif choice == "Save as Word":
#             generate_report_word()

#     # --- Dropdown Button ---
#     dropdown_var = tk.StringVar(value="Generate Report")
#     generate_dropdown = ttk.Menubutton(button_row, textvariable=dropdown_var, width=18)
#     menu = tk.Menu(generate_dropdown, tearoff=0)
#     menu.add_command(label="Save as PDF", command=lambda: save_report_format("Save as PDF"))
#     menu.add_command(label="Save as Excel", command=lambda: save_report_format("Save as Excel"))
#     menu.add_command(label="Save as Word", command=lambda: save_report_format("Save as Word"))
#     generate_dropdown["menu"] = menu
#     generate_dropdown.grid(row=0, column=1, padx=10, ipadx=5, ipady=6)



    # -------- PDF REPORT GENERATION -------- #
    def generate_report_pdf():
        file_name = f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)

        c = pdf_canvas.Canvas(file_path, pagesize=A4)
        width, height = A4
        y = height - 50

        c.setFillColorRGB(0.95, 0.95, 0.95)  
        c.rect(0, height -70, width, 30, stroke=0, fill=1)

        c.setFont("Helvetica-Bold", 16)
        c.setFillColor(colors.HexColor("#1a237e"))
        c.drawCentredString(width / 2, y, " PATIENT REPORT")
        y -= 35

        c.setStrokeColor(colors.lightgrey)
        c.line(40, y, width - 40, y)
        y -= 20

        def draw_section(title, lines):
            nonlocal y
            c.setFont("Helvetica-Bold", 12)
            c.setFillColor(colors.HexColor("#3f51b5"))
            c.drawString(50, y, title)
            y -= 20

            c.setFont("Courier", 11)
            c.setFillColor(colors.black)
            for line in lines:
                c.drawString(60, y, line)
                y -= 16
                if y < 50:
                    c.showPage()
                    y = height - 50

        draw_section(" Company Details", [
            f"Company Name     : {entries['Company Name:'].get()}",
            f"Company Address  : {entries['Company Address:'].get()}",
            f"Customer ID      : {entries['Customer ID:'].get()}",
            f"Mobile No        : {entries['Mobile No:'].get()}",
        ])

        draw_section(" Temperature Analysis", [
            f"A     : {entries['A:'].get()} °C",
            f"B : {entries['B:'].get()} °C",
            f"C    : {entries['C:'].get()} °C",
        ])

        draw_section(" Code Readings", [
            f"D                : {entries['D:'].get()}",
            f"E               : {entries['E:'].get()}",
            f"F                : {entries['F:'].get()}",
    
        ])

        y -= 10
        c.setFillColor(colors.gray)
        c.drawString(50, y, "-" * 70)
        y -= 16
        c.setFont("Helvetica-Oblique", 10)
        c.drawString(50, y, "-- End of Report --")

        c.save()
        messagebox.showinfo("PDF Saved", f"Report saved to:\n{file_path}")


    def generate_report_excel():
        data = {
            "Company Name": [entries['Company Name:'].get()],
            "Company Address": [entries['Company Address:'].get()],
            "Customer ID": [entries['Customer ID:'].get()],
            "Mobile No": [entries['Mobile No:'].get()],
            "A": [entries['A:'].get()],
            "B": [entries['B:'].get()],
            "C": [entries['C:'].get()],
            "D": [entries['D:'].get()],
            "E": [entries['E:'].get()],
            "F": [entries['F:'].get()],
        }

        df = pd.DataFrame(data)
        file_name = f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)
        df.to_excel(file_path, index=False)
        messagebox.showinfo("Excel Saved", f"Report saved to:\n{file_path}")


    def generate_report_word():
        doc = Document()
        doc.add_heading('Patient Report', 0)

        doc.add_heading('Company Details', level=1)
        doc.add_paragraph(f"Company Name: {entries['Company Name:'].get()}")
        doc.add_paragraph(f"Company Address: {entries['Company Address:'].get()}")
        doc.add_paragraph(f"Customer ID: {entries['Customer ID:'].get()}")
        doc.add_paragraph(f"Mobile No: {entries['Mobile No:'].get()}")

        doc.add_heading('Temperature Analysis', level=1)
        doc.add_paragraph(f"A: {entries['A:'].get()} °C")
        doc.add_paragraph(f"B: {entries['B:'].get()} °C")
        doc.add_paragraph(f"C: {entries['C:'].get()} °C")

        doc.add_heading('Code Readings', level=1)
        doc.add_paragraph(f"D: {entries['D:'].get()}")
        doc.add_paragraph(f"E: {entries['E:'].get()}")
        doc.add_paragraph(f"F: {entries['F:'].get()}")

        file_name = f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(os.path.expanduser("~/Downloads"), file_name)
        doc.save(file_path)
        messagebox.showinfo("Word Saved", f"Report saved to:\n{file_path}")


    # # --- Dropdown Function and Button --- #
    # def save_report_format(choice):
    #     if choice == "Save as PDF":
    #         generate_report_pdf()
    #     elif choice == "Save as Excel":
    #         generate_report_excel()
    #     elif choice == "Save as Word":
    #         generate_report_word()

    # dropdown_var = tk.StringVar(value="Generate Report")
    # generate_dropdown = ttk.Menubutton(button_row, textvariable=dropdown_var, width=18)
    # menu = tk.Menu(generate_dropdown, tearoff=0)
    # menu.add_command(label="Save as PDF", command=lambda: save_report_format("Save as PDF"))
    # menu.add_command(label="Save as Excel", command=lambda: save_report_format("Save as Excel"))
    # menu.add_command(label="Save as Word", command=lambda: save_report_format("Save as Word"))
    # generate_dropdown["menu"] = menu
    # generate_dropdown.grid(row=0, column=1, padx=10, ipadx=5, ipady=6)
    
    
    # --- Label Above Dropdown Button ---
    # --- Stylish Dropdown Button with Label as Text ---
    
    
    style = ttk.Style()
    style.configure("Custom.TMenubutton", font=("Segoe UI", 11, "bold"), background="#1a759f",
                    foreground="white", padding=8)
    style.map("Custom.TMenubutton",
            background=[("active", "#166086")],
            foreground=[("active", "white")])

    # --- Dropdown Button ---
    generate_dropdown = ttk.Menubutton(button_row, text="Generate Report", style="Custom.TMenubutton")
    generate_dropdown.grid(row=0, column=1, padx=8, ipadx=6, ipady=2)

    # --- Menu for Format Options ---
    dropdown_menu = tk.Menu(generate_dropdown, tearoff=0, font=("Segoe UI", 10))
    generate_dropdown["menu"] = dropdown_menu
    

    
    # --- Format Saving Function ---
    def save_report_format(choice):
        if choice == "Save as PDF":
            generate_report_pdf()
        elif choice == "Save as Excel":
            generate_report_excel()
        elif choice == "Save as Word":
            generate_report_word()

    # --- Add Menu Commands ---
    dropdown_menu.add_command(label="Save as PDF", command=lambda: save_report_format("Save as PDF"))
    dropdown_menu.add_command(label="Save as Excel", command=lambda: save_report_format("Save as Excel"))
    dropdown_menu.add_command(label="Save as Word", command=lambda: save_report_format("Save as Word"))



def back_to_main():
    global print_frame
    if print_frame:
        print_frame.destroy()
        print_frame = None
    history_frame.pack_forget()
    main_frame.pack(fill="both", expand=True)





# --- Main Window ---
# root = tk.Tk()
root = ctk.CTk()

root.title("Serial Plotter Login")
root.geometry("500x500")
root.configure(bg="#4a4e69")

# ---------------- LOGIN PAGE ----------------
window_width = root.winfo_screenwidth()     
window_height = root.winfo_screenheight()   

# Load and resize image to window
bg_image = Image.open("images/ram.jpg")
bg_image = bg_image.resize((window_width, window_height), Image.LANCZOS)  
bg_photo = ImageTk.PhotoImage(bg_image)

# Create canvas and place background image
canvas = tk.Canvas(root, width=300, height=300)
canvas.pack(fill="both", expand=True)
canvas.create_image(0, 0, image=bg_photo, anchor="nw")





# Create modern-style login frame on top of canvas
frame = tk.Frame(canvas, bg="white", bd=2, relief="ridge", highlightbackground="#ccc", highlightthickness=2)
frame.place(relx=0.5, rely=0.5, anchor="center", width=300, height=300)

heading = tk.Label(frame, text="Serial Data Plotter", font=("Segoe UI", 18, "bold"), bg="white", fg="#22223b")
heading.pack(pady=(20, 10))

username_label = tk.Label(frame, text="Username", font=("Segoe UI", 11,"bold"), bg="white", anchor="w")
username_label.pack(fill="x", padx=40)
# username_entry = tk.Entry(frame, font=("Segoe UI", 11), bd=1, relief=tk.SOLID)
# username_entry.pack(padx=40, pady=(5, 10), ipady=5, fill="x")

username_canvas = tk.Canvas(frame, width=220, height=35, bg="white", bd=0, highlightthickness=0)
username_canvas.pack(padx=40, pady=(5, 10))
username_canvas.create_oval(0, 0, 35, 35, fill="#f0f0f0", outline="#ccc")
username_canvas.create_oval(185, 0, 220, 35, fill="#f0f0f0", outline="#ccc")
username_canvas.create_rectangle(17, 0, 203, 35, fill="#f0f0f0", outline="#ccc")
username_entry = tk.Entry(username_canvas, font=("Segoe UI", 11), bd=0, relief=tk.FLAT, bg="#f0f0f0", highlightthickness=0)
username_entry.place(x=10, y=5, width=200, height=25)

password_label = tk.Label(frame, text="Password", font=("Segoe UI", 11,"bold") ,bg="white", anchor="w")
password_label.pack(fill="x", padx=40)
# password_entry = tk.Entry(frame, font=("Segoe UI", 11), show="*", bd=1, relief=tk.SOLID)
# password_entry.pack(padx=40, pady=(5, 15), ipady=5, fill="x")
password_canvas = tk.Canvas(frame, width=220, height=35, bg="white", bd=0, highlightthickness=0)
password_canvas.pack(padx=40, pady=(5, 15))
password_canvas.create_oval(0, 0, 35, 35, fill="#f0f0f0", outline="#ccc")
password_canvas.create_oval(185, 0, 220, 35, fill="#f0f0f0", outline="#ccc")
password_canvas.create_rectangle(17, 0, 203, 35, fill="#f0f0f0", outline="#ccc")
password_entry = tk.Entry(password_canvas, font=("Segoe UI", 11), show="*", bd=0, relief=tk.FLAT, bg="#f0f0f0", highlightthickness=0)
password_entry.place(x=10, y=5, width=200, height=25)

# login_button = tk.Button(frame, text="Login", font=("Segoe UI", 11, "bold"),
#                          bg="#694a54", fg="white", bd=0, command=login,
#                          cursor="hand2", activebackground="#22223b")
# login_button.pack(pady=(10, 0), ipadx=10, ipady=5)
# --- Rounded Login Button using Canvas ---
def on_login_hover(event):
    login_canvas.itemconfig(login_button_rect, fill="#22223b")

def on_login_leave(event):
    login_canvas.itemconfig(login_button_rect, fill="#694a54")

def on_login_click(event):
    login()

login_canvas = tk.Canvas(frame, width=180, height=40, bg="white", bd=0, highlightthickness=0)
login_canvas.pack(pady=(10, 0))

def create_rounded_rect(canvas, x1, y1, x2, y2, r=20, **kwargs):
    points = [x1+r, y1,
              x2-r, y1,
              x2, y1,
              x2, y1+r,
              x2, y2-r,
              x2, y2,
              x2-r, y2,
              x1+r, y2,
              x1, y2,
              x1, y2-r,
              x1, y1+r,
              x1, y1]
    return canvas.create_polygon(points, smooth=True, **kwargs)

login_button_rect = create_rounded_rect(login_canvas, 0, 0, 180, 40, r=20, fill="#694a54")
login_text = login_canvas.create_text(90, 20, text="Login", fill="white", font=("Segoe UI", 11, "bold"))

login_canvas.bind("<Enter>", on_login_hover)
login_canvas.bind("<Leave>", on_login_leave)
login_canvas.bind("<Button-1>", on_login_click)


# ---------------- MAIN PAGE AFTER LOGIN ----------------
main_frame = tk.Frame(root, bg="#bebebe")
inner_page = tk.Frame(main_frame, bg="#f5f5f5", bd=2, relief="groove")


inner_page.place(relx=0.5, rely=0.5, anchor="center", width=1000, height=700)

main_heading = tk.Label(inner_page, text="Serial Data Monitor", font=("Segoe UI", 22, "bold"), bg="white", fg="#1a1a1a")
main_heading.pack(pady=(40, 20))

# --- Dropdowns ---
dropdown_frame = tk.Frame(inner_page, bg="white")
dropdown_frame.pack(pady=20)

baud_label = tk.Label(dropdown_frame, text="Baud Rate:", font=("Segoe UI", 12), bg="white")
baud_label.pack(side="left", padx=(0, 5))
baud_dropdown = ttk.Combobox(dropdown_frame, values=baud_rates, font=("Segoe UI", 11), state="readonly", width=12)
baud_dropdown.set("9600")
baud_dropdown.pack(side="left", padx=(0, 20))

port_label = tk.Label(dropdown_frame, text="Serial Port:", font=("Segoe UI", 12), bg="white")
port_label.pack(side="left", padx=(0, 5))


port_dropdown = ttk.Combobox(dropdown_frame, values=serial_ports, font=("Segoe UI", 11), state="readonly", width=12)
if serial_ports:
    port_dropdown.set(serial_ports[0])
else:
    port_dropdown.set("No Ports Found")
port_dropdown.pack(side="left")




graph_container = tk.Frame(inner_page, bg="black", width=1500, height=1200, relief="sunken", bd=2)
graph_container.pack(pady=30)

#  DATA TABLE will come here
data_table = tk.Frame(graph_container, bg="black")
data_table.pack(fill="both", expand=True, padx=0, pady=0)

row_widgets = []  #  For tracking and removing old rows

# Placeholder label (optional, can remove if table is showing)
graph_placeholder = tk.Label(graph_container, fg="white", bg="black", font=("Segoe UI", 11))
graph_placeholder.place(relx=0.5, rely=0.5, anchor="center")






def update_table(values):
    global current_index, latest_serial_values
    latest_serial_values = values  
    
    serial_data_history.append(values)


    row_data = "   ".join(values)

    if len(table_rows) < max_rows:
        label = tk.Label(graph_container, text=row_data, fg="white", bg="black", anchor="w", font=("Segoe UI", 11))
        label.pack(fill="x", anchor="w", padx=10, pady=2)
        table_rows.append(label)
    else:
        table_rows[current_index].config(text=row_data)
        current_index = (current_index + 1) % max_rows



table_rows = []        # stores Label rows
max_rows = 10          # max 10 rows visible
current_index = 0      # points to row to overwrite next
graph_placeholder.place(relx=0.5, rely=0.5, anchor="center")


def read_serial_data():
    global ser
    while running:
        try:
            line = ser.readline().decode('utf-8').strip()
            if line:
                values = line.split(',')
                if len(values) == 6:
                    update_table(values)
        except Exception as e:
            print("Serial Read Error:", e)
        time.sleep(0.1)



style = ttk.Style()
style.theme_use('default')

style.configure("Rounded.TButton",
    font=("Segoe UI", 12, "bold"),
    foreground="white",
    background="#4a4e69",  # default background (can override per button)
    padding=10,
    borderwidth=0
)
style.map("Rounded.TButton",
    background=[("active", "#22223b")],
    foreground=[("pressed", "white")],
    relief=[("pressed", "flat")]
)

style.configure("Start.TButton", background="#4361ee")
style.map("Start.TButton", background=[("active", "#364fc7")])

style.configure("Stop.TButton", background="#ef476f")
style.map("Stop.TButton", background=[("active", "#d6336c")])

style.configure("Print.TButton", background="#06d6a0")
style.map("Print.TButton", background=[("active", "#38b000")])

style.configure("History.TButton", background="#adb5bd")
style.map("History.TButton", background=[("active", "#6c757d")])

bottom_button_frame = tk.Frame(inner_page, bg="#f5f5f5")
bottom_button_frame.pack(side="bottom", pady=20)



# Start Button - Indigo Blue
start_btn = ttk.Button(bottom_button_frame, text="Start", style="Rounded.TButton", command=start_serial)
start_btn.pack(side="left", padx=10, ipadx=10, ipady=6)
start_btn.configure(style="Start.TButton")

# Stop Button - Coral Red
stop_btn = ttk.Button(bottom_button_frame, text="Stop", style="Rounded.TButton", command=stop_serial)
stop_btn.pack(side="left", padx=10, ipadx=10, ipady=6)
stop_btn.configure(style="Stop.TButton")

# Print Button - Mint Green
print_btn = ttk.Button(bottom_button_frame, text="Print", style="Rounded.TButton", command=show_print_page)
print_btn.pack(side="left", padx=10, ipadx=10, ipady=6)
print_btn.configure(style="Print.TButton")

# History Button - Soft Gray
history_btn = ttk.Button(bottom_button_frame, text="History", style="Rounded.TButton", command=show_history_page)
history_btn.pack(side="left", padx=10, ipadx=10, ipady=6)
history_btn.configure(style="History.TButton")





# --- History Page ---
history_frame = tk.Frame(root, bg="#e3f2fd")  # soft blue background

# 1. Header Section
history_heading = tk.Label(history_frame, text=" History Overview", 
                           font=("Segoe UI", 24, "bold"), bg="#e3f2fd", fg="#0d47a1")
history_heading.pack(pady=(30, 10))

separator1 = tk.Frame(history_frame, bg="#90caf9", height=2)
separator1.pack(fill="x", padx=100)



root.mainloop()
